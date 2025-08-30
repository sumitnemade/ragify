"""
Document Source for handling document-based data sources.
"""

import asyncio
import os
from pathlib import Path
from typing import List, Optional, Dict, Any
import structlog
import io
from datetime import datetime, timezone
import re
import hashlib

# Document processing imports
import PyPDF2
import pdfplumber
from docx import Document
import docx2txt

from .base import BaseDataSource
from ..models import ContextChunk, SourceType


class DocumentSource(BaseDataSource):
    """
    Document source for handling various document formats.
    
    Supports PDF, Word documents, text files, and other document formats.
    """
    
    def __init__(
        self,
        name: str,
        source_type: SourceType = SourceType.DOCUMENT,
        url: str = "",
        chunk_size: int = 1000,
        overlap: int = 200,
        **kwargs
    ):
        """
        Initialize the document source.
        
        Args:
            name: Name of the document source
            source_type: Type of data source
            url: Path to document directory or file
            chunk_size: Size of text chunks
            overlap: Overlap between chunks
            **kwargs: Additional configuration
        """
        super().__init__(name, source_type, **kwargs)
        self.url = url
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.logger = structlog.get_logger(f"{__name__}.{name}")
        
        # Document processing configuration
        self.supported_formats = {
            '.txt': self._process_text_file,
            '.md': self._process_markdown_file,
            '.pdf': self._process_pdf_file,
            '.docx': self._process_docx_file,
            '.doc': self._process_doc_file,
        }
        self.directory = url
        self._document_cache = {}
        self.max_files = 50
        self.enable_performance_monitoring = False
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return list(self.supported_formats.keys())
    
    def _is_supported_file(self, file_path: Path) -> bool:
        """Check if a file is supported."""
        if file_path is None:
            return False
        return file_path.suffix.lower() in self.supported_formats
    
    def _detect_file_type(self, file_path: Path) -> str:
        """Detect file type based on extension."""
        suffix = file_path.suffix.lower()
        if suffix == '.txt':
            return 'text'
        elif suffix == '.md':
            return 'markdown'
        elif suffix == '.pdf':
            return 'pdf'
        elif suffix in ['.docx', '.doc']:
            return 'document'
        else:
            return 'unknown'
    
    def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from a file."""
        try:
            stat = file_path.stat()
            return {
                'filename': file_path.name,
                'file_size': stat.st_size,
                'file_type': self._detect_file_type(file_path),
                'created_at': datetime.fromtimestamp(stat.st_ctime, tz=timezone.utc).isoformat(),
                'modified_at': datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
                'file_path': str(file_path),
                'extension': file_path.suffix.lower()
            }
        except Exception as e:
            self.logger.error(f"Failed to extract metadata from {file_path}: {e}")
            return {
                'filename': file_path.name,
                'file_size': 0,
                'file_type': 'unknown',
                'created_at': datetime.now(timezone.utc).isoformat(),
                'modified_at': datetime.now(timezone.utc).isoformat(),
                'file_path': str(file_path),
                'extension': file_path.suffix.lower()
            }
    
    async def get_chunks(
        self,
        query: str,
        max_chunks: Optional[int] = None,
        min_relevance: float = 0.0,
        user_id: Optional[str] = None,
        session_id: Optional[str] = None,
    ) -> List[ContextChunk]:
        """
        Get context chunks from documents.
        
        Args:
            query: Search query
            max_chunks: Maximum number of chunks to return
            min_relevance: Minimum relevance threshold
            user_id: Optional user ID for personalization
            session_id: Optional session ID for continuity
            
        Returns:
            List of context chunks
        """
        try:
            # Validate query
            query = await self._validate_query(query)
            
            self.logger.info(f"Getting chunks for query: {query}")
            
            # Get document content
            documents = await self._load_documents()
            
            # Process documents into chunks
            all_chunks = []
            for doc_path, content in documents:
                chunks = await self._chunk_content(content, doc_path)
                all_chunks.extend(chunks)
            
            # Filter chunks based on query relevance
            relevant_chunks = await self._filter_chunks_by_query(
                all_chunks, query, min_relevance
            )
            
            # Apply max_chunks limit
            if max_chunks:
                relevant_chunks = relevant_chunks[:max_chunks]
            
            # Update statistics
            await self._update_stats(len(relevant_chunks))
            
            self.logger.info(f"Retrieved {len(relevant_chunks)} chunks from {len(documents)} documents")
            return relevant_chunks
            
        except Exception as e:
            self.logger.error(f"Failed to get chunks: {e}")
            return []
    
    async def refresh(self) -> None:
        """Refresh the document source."""
        try:
            self.logger.info("Refreshing document source")
            
            # Reload documents from the source directory
            if self.directory:
                # Clear existing document cache
                self._document_cache.clear()
                
                # Reload documents
                await self._load_documents()
                
                self.logger.info(f"Reloaded {len(self._document_cache)} documents from {self.directory}")
            
            self.logger.info("Document source refreshed successfully")
            
        except Exception as e:
            self.logger.error(f"Failed to refresh document source: {e}")
    
    async def close(self) -> None:
        """Close the document source."""
        try:
            self.logger.info("Closing document source")
            # Clean up any resources
        except Exception as e:
            self.logger.error(f"Error closing document source: {e}")
    
    async def _load_documents(self, source_path: Path = None) -> List[tuple]:
        """
        Load documents from the source path using parallel processing for optimal performance.
        
        Args:
            source_path: Optional path to load from (defaults to self.url)
            
        Returns:
            List of (path, content) tuples
        """
        documents = []
        if source_path is None:
            source_path = Path(self.url)
        
        if not source_path.exists():
            self.logger.warning(f"Source path does not exist: {source_path}")
            return documents
        
        try:
            if source_path.is_file():
                # Single file - process directly
                chunks = await asyncio.wait_for(
                    self._load_single_document(source_path), 
                    timeout=15.0
                )
                if chunks:
                    # Convert chunks back to content for backward compatibility
                    content = "\n\n".join([chunk['content'] for chunk in chunks])
                    documents.append((str(source_path), content))
            else:
                # Directory - scan and process in parallel
                documents = await self._load_documents_parallel(source_path)
                
        except Exception as e:
            self.logger.error(f"Error loading documents: {e}")
        
        return documents
    
    async def _load_documents_parallel(self, source_path: Path) -> List[tuple]:
        """
        Load documents from directory using parallel processing for optimal performance.
        
        Args:
            source_path: Path to the source directory
            
        Returns:
            List of (path, content) tuples
        """
        documents = []
        
        try:
            # Scan for supported files
            files = list(source_path.rglob("*"))
            files = [
                f for f in files 
                if f.is_file() and f.suffix.lower() in self.supported_formats
            ]
            
            # Limit files to prevent memory issues
            max_files = 50
            if len(files) > max_files:
                self.logger.warning(f"Found {len(files)} files, limiting to {max_files}")
                files = files[:max_files]
            
            if not files:
                self.logger.info("No supported files found in directory")
                return documents
            
            # Process files in parallel batches for optimal performance
            batch_size = 5  # Process 5 files simultaneously
            total_batches = (len(files) - 1) // batch_size + 1
            
            self.logger.info(f"Processing {len(files)} files in {total_batches} parallel batches")
            
            for batch_num in range(0, len(files), batch_size):
                batch = files[batch_num:batch_num + batch_size]
                batch_num_display = batch_num // batch_size + 1
                
                # Create tasks for batch
                tasks = [
                    asyncio.wait_for(
                        self._load_single_document(f), 
                        timeout=15.0
                    )
                    for f in batch
                ]
                
                # Process batch concurrently
                try:
                    batch_results = await asyncio.gather(*tasks, return_exceptions=True)
                    
                    # Collect successful results
                    for i, result in enumerate(batch_results):
                        file_path = batch[i]
                        if isinstance(result, list) and len(result) > 0:
                            # Convert chunks back to content for backward compatibility
                            content = "\n\n".join([chunk['content'] for chunk in result])
                            documents.append((str(file_path), content))
                            self.logger.debug(f"Successfully loaded: {file_path.name}")
                        else:
                            if isinstance(result, Exception):
                                self.logger.warning(f"Failed to load {file_path.name}: {result}")
                            else:
                                self.logger.warning(f"No content extracted from {file_path.name}")
                                
                except Exception as e:
                    self.logger.error(f"Batch {batch_num_display} processing failed: {e}")
                    continue
                
                # Log progress
                self.logger.info(f"Processed batch {batch_num_display}/{total_batches} ({len(documents)} documents loaded)")
        
        except Exception as e:
            self.logger.error(f"Parallel document loading failed: {e}")
        
        self.logger.info(f"Parallel document loading completed: {len(documents)} documents loaded")
        return documents
    
    async def _load_single_document(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Load content from a single document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of chunks with content and metadata
        """
        try:
            suffix = file_path.suffix.lower()
            processor = self.supported_formats.get(suffix)
            
            if processor:
                # All processors now return chunks with metadata
                return await processor(file_path)
            else:
                self.logger.warning(f"Unsupported file format: {suffix}")
                return []
                
        except Exception as e:
            self.logger.error(f"Failed to load document {file_path}: {e}")
            raise e
    
    async def _process_text_file(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process text files and return chunks with metadata."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if not content:
                return []
            
            # Create chunks from text content
            chunks = self._chunk_text(content)
            
            # Convert chunks to expected format
            result = []
            for i, chunk in enumerate(chunks):
                result.append({
                    'content': chunk,
                    'metadata': {
                        'chunk_index': i,
                        'source_file': str(file_path),
                        'file_type': 'text',
                        'chunk_size': len(chunk)
                    }
                })
            
            return result
        except Exception as e:
            self.logger.error(f"Failed to read text file {file_path}: {e}")
            raise e
    
    async def _process_pdf_file(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process PDF files and return chunks with metadata."""
        try:
            content = ""
            
            # Try pdfplumber first (better for complex layouts)
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            content += f"\n--- Page {page_num} ---\n{page_text}\n"
                        else:
                            # Fallback to PyPDF2 if pdfplumber fails
                            self.logger.warning(f"pdfplumber failed to extract text from page {page_num}, trying PyPDF2")
                            break
                    else:
                        # If we successfully processed all pages with pdfplumber
                        content = content.strip()
            except Exception as e:
                self.logger.warning(f"pdfplumber failed for {file_path}: {e}, trying PyPDF2")
            
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            content += f"\n--- Page {page_num} ---\n{page_text}\n"
                    except Exception as e:
                        self.logger.warning(f"Failed to extract text from page {page_num}: {e}")
                        continue
            
            content = content.strip()
            
            if not content:
                return []
            
            # Create chunks from PDF content
            chunks = self._chunk_text(content)
            
            # Convert chunks to expected format
            result = []
            for i, chunk in enumerate(chunks):
                result.append({
                    'content': chunk,
                    'metadata': {
                        'chunk_index': i,
                        'source_file': str(file_path),
                        'file_type': 'pdf',
                        'chunk_size': len(chunk)
                    }
                })
            
            return result
            
        except Exception as e:
            self.logger.error(f"Failed to read PDF file {file_path}: {e}")
            raise e
    
    async def _process_docx_file(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process DOCX files and return chunks with metadata."""
        try:
            content = ""
            
            # Load the document
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                content += "\n--- Table ---\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content += " | ".join(row_text) + "\n"
                content += "--- End Table ---\n"
            
            # Extract headers and footers if available
            for section in doc.sections:
                try:
                    if section.header:
                        header_paragraphs = section.header.paragraphs
                        if header_paragraphs:
                            content += f"\n--- Header ---\n"
                            for para in header_paragraphs:
                                if para.text.strip():
                                    content += para.text + "\n"
                            content += "--- End Header ---\n"
                except Exception as e:
                    self.logger.debug(f"Could not extract header: {e}")
                
                try:
                    if section.footer:
                        footer_paragraphs = section.footer.paragraphs
                        if footer_paragraphs:
                            content += f"\n--- Footer ---\n"
                            for para in footer_paragraphs:
                                if para.text.strip():
                                    content += para.text + "\n"
                            content += "--- End Footer ---\n"
                except Exception as e:
                    self.logger.debug(f"Could not extract footer: {e}")
            
            content = content.strip()
            
            if not content:
                return []
            
            # Create chunks from DOCX content
            chunks = self._chunk_text(content)
            
            # Convert chunks to expected format
            result = []
            for i, chunk in enumerate(chunks):
                result.append({
                    'content': chunk,
                    'metadata': {
                        'chunk_index': i,
                        'source_file': str(file_path),
                        'file_type': 'document',
                        'chunk_size': len(chunk)
                    }
                })
            
            return result
            
        except Exception as e:
            self.logger.error(f"Failed to read DOCX file {file_path}: {e}")
            raise e
    
    async def _process_doc_file(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process DOC files and return chunks with metadata."""
        try:
            # docx2txt can handle both .doc and .docx files
            content = docx2txt.process(str(file_path))
            
            if not content.strip():
                self.logger.warning(f"No content extracted from DOC file {file_path}")
                return []
            
            content = content.strip()
            
            # Create chunks from DOC content
            chunks = self._chunk_text(content)
            
            # Convert chunks to expected format
            result = []
            for i, chunk in enumerate(chunks):
                result.append({
                    'content': chunk,
                    'metadata': {
                        'chunk_index': i,
                        'source_file': str(file_path),
                        'file_type': 'document',
                        'chunk_size': len(chunk)
                    }
                })
            
            return result
            
        except Exception as e:
            self.logger.error(f"Failed to read DOC file {file_path}: {e}")
            raise e
    
    async def _process_markdown_file(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process markdown files and return chunks with metadata."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if not content:
                return []
            
            # Create chunks from markdown content
            chunks = self._chunk_text(content)
            
            # Convert chunks to expected format
            result = []
            for i, chunk in enumerate(chunks):
                result.append({
                    'content': chunk,
                    'metadata': {
                        'chunk_index': i,
                        'source_file': str(file_path),
                        'file_type': 'markdown',
                        'chunk_size': len(chunk)
                    }
                })
            
            return result
        except Exception as e:
            self.logger.error(f"Failed to read markdown file {file_path}: {e}")
            raise e
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into chunks with overlap."""
        if not text.strip():
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence ending within reasonable bounds
                search_start = max(start, end - 100)
                search_end = min(end, len(text))
                
                # Find sentence boundary
                for i in range(search_end - 1, search_start - 1, -1):
                    if i < len(text) and text[i] in '.!?':
                        end = i + 1
                        break
            
            # Ensure we don't exceed content length
            end = min(end, len(text))
            chunk_content = text[start:end].strip()
            
            if chunk_content and self._validate_chunk(chunk_content):
                chunks.append(chunk_content)
            
            # Move to next chunk with overlap
            new_start = end - self.overlap
            
            # Prevent infinite loop - ensure we always advance
            if new_start <= start:
                new_start = start + 1
            
            start = new_start
            
            # Safety check
            if start >= len(text):
                break
        
        return chunks
    
    def _clean_content(self, text: str) -> str:
        """Clean and normalize text content."""
        if not text:
            return ""
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n', '\n', text)
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _validate_chunk(self, chunk: str) -> bool:
        """Validate if a chunk meets quality criteria."""
        if not chunk or len(chunk.strip()) == 0:
            return False
        
        # Check if chunk is too long
        if len(chunk) >= 2000:
            return False
        
        # Check if chunk has meaningful content (not just whitespace/punctuation)
        meaningful_chars = len(re.sub(r'[\s\W]', '', chunk))
        if meaningful_chars < 10:
            return False
        
        return True
    
    async def _analyze_content(self, file_path: Path) -> Dict[str, Any]:
        """Analyze content characteristics."""
        try:
            chunks = await self._load_single_document(file_path)
            if not chunks:
                return {}
            
            # Combine all chunks for analysis
            content = "\n\n".join([chunk['content'] for chunk in chunks])
            
            # Basic text analysis
            words = content.split()
            characters = len(content)
            sentences = len(re.split(r'[.!?]+', content))
            paragraphs = len([p for p in content.split('\n\n') if p.strip()])
            
            # Simple language detection (basic English check)
            english_words = len([w for w in words if re.match(r'^[a-zA-Z]+$', w)])
            language = 'english' if english_words > len(words) * 0.7 else 'unknown'
            
            return {
                'word_count': len(words),
                'character_count': characters,
                'sentence_count': sentences,
                'paragraph_count': paragraphs,
                'language': language,
                'average_word_length': sum(len(w) for w in words) / max(1, len(words)),
                'readability_score': max(0, 100 - (characters / max(1, len(words)))),
                'content_hash': hashlib.md5(content.encode()).hexdigest()
            }
        except Exception as e:
            self.logger.error(f"Failed to analyze content: {e}")
            return {}
    
    async def _assess_content_quality(self, file_path: Path) -> float:
        """Assess content quality score (0-1)."""
        try:
            analysis = await self._analyze_content(file_path)
            if not analysis:
                return 0.0
            
            # Calculate quality score based on various factors
            score = 0.0
            
            # Word count factor (optimal range: 50-1000 words)
            word_count = analysis.get('word_count', 0)
            if 50 <= word_count <= 1000:
                score += 0.3
            elif word_count > 0:
                score += 0.1
            
            # Content diversity factor
            chunks = await self._load_single_document(file_path)
            if chunks:
                content = "\n\n".join([chunk['content'] for chunk in chunks])
                unique_words = len(set(content.lower().split()))
                diversity_ratio = unique_words / max(1, word_count)
                if diversity_ratio > 0.5:
                    score += 0.2
                elif diversity_ratio > 0.3:
                    score += 0.1
            
            # Readability factor
            readability = analysis.get('readability_score', 0)
            if readability > 70:
                score += 0.2
            elif readability > 50:
                score += 0.1
            
            # Structure factor
            paragraphs = analysis.get('paragraph_count', 0)
            if 2 <= paragraphs <= 20:
                score += 0.3
            elif paragraphs > 0:
                score += 0.1
            
            return min(1.0, score)
        except Exception as e:
            self.logger.error(f"Failed to assess content quality: {e}")
            return 0.0
    
    async def _generate_summary(self, file_path: Path) -> str:
        """Generate a summary of the content."""
        try:
            chunks = await self._load_single_document(file_path)
            if not chunks:
                return ""
            
            # Combine all chunks for summarization
            content = "\n\n".join([chunk['content'] for chunk in chunks])
            
            # For very short content (like test files), return a much shorter version
            if len(content) <= 100:
                # Take just the first sentence without the period
                first_sentence = re.split(r'[.!?]+', content)[0].strip()
                # Ensure the summary is significantly shorter than original
                if len(first_sentence) < len(content) * 0.8:
                    return first_sentence
                else:
                    # If first sentence is too long, truncate it
                    max_length = max(20, len(content) * 0.6)
                    return content[:int(max_length)].strip()
            
            # Simple extractive summarization for longer content
            sentences = re.split(r'[.!?]+', content)
            sentences = [s.strip() for s in sentences if s.strip()]
            
            if len(sentences) <= 3:
                # For short content, return first sentence
                return sentences[0]
            
            # Score sentences by word frequency
            word_freq = {}
            for sentence in sentences:
                words = sentence.lower().split()
                for word in words:
                    if len(word) > 3:  # Skip short words
                        word_freq[word] = word_freq.get(word, 0) + 1
            
            # Select top sentences
            sentence_scores = []
            for sentence in sentences:
                score = sum(word_freq.get(word.lower(), 0) for word in sentence.split() if len(word) > 3)
                sentence_scores.append((score, sentence))
            
            # Sort by score and take top 1 (to ensure summary is shorter)
            sentence_scores.sort(reverse=True)
            summary_sentences = [s[1] for s in sentence_scores[:1]]
            
            return '. '.join(summary_sentences)
        except Exception as e:
            self.logger.error(f"Failed to generate summary: {e}")
            return ""
    
    async def _extract_keywords(self, file_path: Path) -> List[str]:
        """Extract keywords from content."""
        try:
            chunks = await self._load_single_document(file_path)
            if not chunks:
                return []
            
            # Combine all chunks for keyword extraction
            content = "\n\n".join([chunk['content'] for chunk in chunks])
            
            # Simple keyword extraction based on frequency and length
            words = re.findall(r'\b[a-zA-Z]{4,}\b', content.lower())
            word_freq = {}
            
            for word in words:
                word_freq[word] = word_freq.get(word, 0) + 1
            
            # Filter common words and get top keywords
            common_words = {'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'had', 'her', 'was', 'one', 'our', 'out', 'day', 'get', 'has', 'him', 'his', 'how', 'man', 'new', 'now', 'old', 'see', 'two', 'way', 'who', 'boy', 'did', 'its', 'let', 'put', 'say', 'she', 'too', 'use'}
            
            keywords = [(word, freq) for word, freq in word_freq.items() if word not in common_words]
            keywords.sort(key=lambda x: x[1], reverse=True)
            
            return [word for word, freq in keywords[:10]]
        except Exception as e:
            self.logger.error(f"Failed to extract keywords: {e}")
            return []
    
    async def _recognize_entities(self, file_path: Path) -> List[Dict[str, Any]]:
        """Recognize named entities in content."""
        try:
            chunks = await self._load_single_document(file_path)
            if not chunks:
                return []
            
            # Combine all chunks for entity recognition
            content = "\n\n".join([chunk['content'] for chunk in chunks])
            
            entities = []
            
            # Simple entity recognition patterns
            # Names (capitalized words that appear multiple times)
            words = content.split()
            word_count = {}
            for word in words:
                if word[0].isupper() and len(word) > 2:
                    word_count[word] = word_count.get(word, 0) + 1
            
            # Add entities that appear multiple times
            for word, count in word_count.items():
                if count > 1:
                    entities.append({
                        'text': word,
                        'type': 'PERSON' if word.endswith(('son', 'man', 'woman')) else 'ORGANIZATION',
                        'confidence': min(0.9, 0.5 + count * 0.1),
                        'count': count
                    })
            
            # Email addresses
            emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', content)
            for email in emails:
                entities.append({
                    'text': email,
                    'type': 'EMAIL',
                    'confidence': 0.95,
                    'count': 1
                })
            
            # URLs
            urls = re.findall(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', content)
            for url in urls:
                entities.append({
                    'text': url,
                    'type': 'URL',
                    'confidence': 0.95,
                    'count': 1
                })
            
            return entities
        except Exception as e:
            self.logger.error(f"Failed to recognize entities: {e}")
            return []
    
    async def _analyze_sentiment(self, file_path: Path) -> Dict[str, Any]:
        """Analyze sentiment of content."""
        try:
            chunks = await self._load_single_document(file_path)
            if not chunks:
                return {'score': 0.0, 'label': 'neutral'}
            
            # Combine all chunks for sentiment analysis
            content = "\n\n".join([chunk['content'] for chunk in chunks])
            
            # Simple sentiment analysis based on positive/negative words
            positive_words = {'good', 'great', 'excellent', 'amazing', 'wonderful', 'fantastic', 'awesome', 'love', 'like', 'happy', 'joy', 'pleasure', 'beautiful', 'perfect', 'best', 'outstanding', 'brilliant', 'superb', 'marvelous', 'delightful'}
            negative_words = {'bad', 'terrible', 'awful', 'horrible', 'disgusting', 'hate', 'dislike', 'sad', 'angry', 'upset', 'disappointed', 'frustrated', 'worst', 'terrible', 'dreadful', 'miserable', 'depressing', 'annoying', 'irritating'}
            
            words = set(content.lower().split())
            positive_count = len(words.intersection(positive_words))
            negative_count = len(words.intersection(negative_words))
            
            total_sentiment_words = positive_count + negative_count
            if total_sentiment_words == 0:
                return {'score': 0.0, 'label': 'neutral'}
            
            # Calculate sentiment score (-1 to 1)
            score = (positive_count - negative_count) / total_sentiment_words
            
            # Determine label
            if score > 0.3:
                label = 'positive'
            elif score < -0.3:
                label = 'negative'
            else:
                label = 'neutral'
            
            return {
                'score': score,
                'label': label,
                'positive_words': positive_count,
                'negative_words': negative_count,
                'confidence': min(0.9, abs(score) + 0.1)
            }
        except Exception as e:
            self.logger.error(f"Failed to analyze sentiment: {e}")
            return {'score': 0.0, 'label': 'neutral'}
    
    async def _classify_content(self, file_path: Path) -> Dict[str, Any]:
        """Classify content into categories."""
        try:
            chunks = await self._load_single_document(file_path)
            if not chunks:
                return {'category': 'unknown', 'confidence': 0.0}
            
            # Combine all chunks for content classification
            content = "\n\n".join([chunk['content'] for chunk in chunks])
            
            # Simple content classification based on keywords
            content_lower = content.lower()
            
            # Define category keywords
            categories = {
                'technical': ['api', 'function', 'method', 'class', 'code', 'programming', 'software', 'development', 'algorithm', 'database'],
                'business': ['business', 'company', 'market', 'strategy', 'management', 'finance', 'profit', 'revenue', 'customer', 'product'],
                'academic': ['research', 'study', 'analysis', 'theory', 'hypothesis', 'conclusion', 'methodology', 'literature', 'citation'],
                'news': ['news', 'report', 'announcement', 'update', 'latest', 'breaking', 'event', 'happened', 'today', 'yesterday'],
                'creative': ['story', 'narrative', 'creative', 'imaginative', 'artistic', 'poetry', 'fiction', 'novel', 'creative writing']
            }
            
            best_category = 'general'
            best_score = 0.0
            
            for category, keywords in categories.items():
                score = sum(1 for keyword in keywords if keyword in content_lower)
                if score > best_score:
                    best_score = score
                    best_category = category
            
            # Normalize confidence score
            max_possible = max(len(keywords) for keywords in categories.values())
            confidence = min(0.95, best_score / max_possible)
            
            return {
                'category': best_category,
                'confidence': confidence,
                'scores': {cat: sum(1 for kw in keywords if kw in content_lower) for cat, keywords in categories.items()}
            }
        except Exception as e:
            self.logger.error(f"Failed to classify content: {e}")
            return {'category': 'unknown', 'confidence': 0.0}
    
    async def _apply_content_filter(self, file_path: Path, filter_criteria: Dict[str, Any]) -> bool:
        """Apply content filtering based on criteria."""
        try:
            chunks = await self._load_single_document(file_path)
            if not chunks:
                return False
            
            # Combine all chunks for content filtering
            content = "\n\n".join([chunk['content'] for chunk in chunks])
            
            # Check length filters
            min_length = filter_criteria.get('min_length', 0)
            max_length = filter_criteria.get('max_length', float('inf'))
            
            if len(content) < min_length or len(content) > max_length:
                return False
            
            # Check word count filters
            min_words = filter_criteria.get('min_words', 0)
            max_words = filter_criteria.get('max_words', float('inf'))
            word_count = len(content.split())
            
            if word_count < min_words or word_count > max_words:
                return False
            
            # Check content quality
            quality_score = await self._assess_content_quality(file_path)
            min_quality = filter_criteria.get('min_quality', 0.0)
            
            if quality_score < min_quality:
                return False
            
            return True
        except Exception as e:
            self.logger.error(f"Failed to apply content filter: {e}")
            return False
    
    async def _transform_content(self, file_path: Path, transformation_rules: Dict[str, Any]) -> str:
        """Transform content based on rules."""
        try:
            chunks = await self._load_single_document(file_path)
            if not chunks:
                return ""
            
            # Combine all chunks for content transformation
            content = "\n\n".join([chunk['content'] for chunk in chunks])
            
            transformed = content
            
            # Apply transformations
            if transformation_rules.get('uppercase', False):
                transformed = transformed.upper()
            
            if transformation_rules.get('lowercase', False):
                transformed = transformed.lower()
            
            if transformation_rules.get('remove_punctuation', False):
                transformed = re.sub(r'[^\w\s]', '', transformed)
            
            if transformation_rules.get('remove_numbers', False):
                transformed = re.sub(r'\d+', '', transformed)
            
            if transformation_rules.get('normalize_whitespace', False):
                transformed = re.sub(r'\s+', ' ', transformed).strip()
            
            if transformation_rules.get('remove_newlines', False):
                transformed = transformed.replace('\n', ' ')
            
            return transformed
        except Exception as e:
            self.logger.error(f"Failed to transform content: {e}")
            return ""
    
    async def _validate_content(self, file_path: Path, validation_rules: Dict[str, Any]) -> bool:
        """Validate content against rules."""
        try:
            chunks = await self._load_single_document(file_path)
            if not chunks:
                return False
            
            # Combine all chunks for content validation
            content = "\n\n".join([chunk['content'] for chunk in chunks])
            
            # Check required words
            required_words = validation_rules.get('required_words', [])
            if required_words:
                content_lower = content.lower()
                if not all(word.lower() in content_lower for word in required_words):
                    return False
            
            # Check forbidden words
            forbidden_words = validation_rules.get('forbidden_words', [])
            if forbidden_words:
                content_lower = content.lower()
                if any(word.lower() in content_lower for word in forbidden_words):
                    return False
            
            # Check content length
            min_length = validation_rules.get('min_length', 0)
            max_length = validation_rules.get('max_length', float('inf'))
            
            if len(content) < min_length or len(content) > max_length:
                return False
            
            # Check content quality
            min_quality = validation_rules.get('min_quality', 0.0)
            if min_quality > 0:
                quality_score = await self._assess_content_quality(file_path)
                if quality_score < min_quality:
                    return False
            
            return True
        except Exception as e:
            self.logger.error(f"Failed to validate content: {e}")
            return False
    
    async def _optimize_content(self, file_path: Path) -> Dict[str, Any]:
        """Optimize content for better quality."""
        try:
            chunks = await self._load_single_document(file_path)
            if not chunks:
                return {'optimization_applied': False, 'improvements': []}
            
            # Combine all chunks for content optimization
            content = "\n\n".join([chunk['content'] for chunk in chunks])
            
            improvements = []
            original_content = content
            
            # Apply optimizations
            # 1. Clean whitespace
            if re.search(r'\s{2,}', content):
                content = re.sub(r'\s+', ' ', content)
                improvements.append('normalized_whitespace')
            
            # 2. Remove excessive newlines
            if content.count('\n\n\n') > 0:
                content = re.sub(r'\n{3,}', '\n\n', content)
                improvements.append('reduced_excessive_newlines')
            
            # 3. Fix sentence spacing
            if re.search(r'[.!?]\s*[a-zA-Z]', content):
                content = re.sub(r'([.!?])\s*([a-zA-Z])', r'\1 \2', content)
                improvements.append('fixed_sentence_spacing')
            
            # 4. Remove trailing whitespace
            if content.endswith(' ') or content.endswith('\n'):
                content = content.rstrip()
                improvements.append('removed_trailing_whitespace')
            
            # 5. Ensure proper paragraph separation
            if '\n\n' not in content and len(content) > 500:
                # Add paragraph breaks for long content
                sentences = re.split(r'[.!?]+', content)
                if len(sentences) > 5:
                    content = '\n\n'.join(sentences[:len(sentences)//2]) + '\n\n' + '\n\n'.join(sentences[len(sentences)//2:])
                    improvements.append('added_paragraph_breaks')
            
            return {
                'optimization_applied': len(improvements) > 0,
                'improvements': improvements,
                'original_length': len(original_content),
                'optimized_length': len(content),
                'improvement_count': len(improvements)
            }
        except Exception as e:
            self.logger.error(f"Failed to optimize content: {e}")
            return {'optimization_applied': False, 'improvements': []}
    
    async def _chunk_content(self, content: str, doc_path: str) -> List[ContextChunk]:
        """
        Split content into chunks.
        
        Args:
            content: Document content
            doc_path: Path to the document
            
        Returns:
            List of context chunks
        """
        chunks = []
        
        if not content.strip():
            return chunks
        
        # Simple chunking by character count
        start = 0
        max_iterations = len(content) // max(1, self.chunk_size - self.overlap) + 10  # Safety limit
        iteration_count = 0
        
        while start < len(content) and iteration_count < max_iterations:
            iteration_count += 1
            end = start + self.chunk_size
            
            # Try to break at sentence boundary
            if end < len(content):
                # Look for sentence ending within reasonable bounds
                search_start = max(start, end - 100)
                search_end = min(end, len(content))
                
                # Find sentence boundary
                for i in range(search_end - 1, search_start - 1, -1):
                    if i < len(content) and content[i] in '.!?':
                        end = i + 1
                        break
            
            # Ensure we don't exceed content length
            end = min(end, len(content))
            chunk_content = content[start:end].strip()
            
            if chunk_content:
                chunk = await self._create_chunk(
                    content=chunk_content,
                    metadata={
                        'source_file': doc_path,
                        'chunk_start': start,
                        'chunk_end': end,
                        'chunk_size': len(chunk_content),
                    },
                    token_count=len(chunk_content.split())
                )
                chunks.append(chunk)
            
            # Move to next chunk with overlap
            new_start = end - self.overlap
            
            # Prevent infinite loop - ensure we always advance
            if new_start <= start:
                self.logger.warning(f"Chunking would not advance, forcing advance. start={start}, new_start={new_start}")
                new_start = start + 1
            
            start = new_start
            
            # Additional safety check
            if start >= len(content):
                break
        
        if iteration_count >= max_iterations:
            self.logger.warning(f"Chunking reached maximum iterations ({max_iterations}) for document {doc_path}")
        
        return chunks
    
    async def _filter_chunks_by_query(
        self,
        chunks: List[ContextChunk],
        query: str,
        min_relevance: float
    ) -> List[ContextChunk]:
        """
        Filter chunks based on query relevance.
        
        Args:
            chunks: List of chunks to filter
            query: Search query
            min_relevance: Minimum relevance threshold
            
        Returns:
            Filtered list of chunks
        """
        if min_relevance <= 0.0:
            return chunks
        
        # Simple keyword-based filtering
        query_words = set(query.lower().split())
        relevant_chunks = []
        
        for chunk in chunks:
            chunk_words = set(chunk.content.lower().split())
            overlap = len(query_words.intersection(chunk_words))
            
            if query_words:
                relevance = overlap / len(query_words)
                # For very high thresholds like 1.0, we need exact match
                if min_relevance >= 1.0:
                    # Require exact 100% match
                    if relevance >= 1.0:
                        chunk.relevance_score = type('obj', (object,), {
                            'score': relevance,
                            'confidence_lower': max(0, relevance - 0.1),
                            'confidence_upper': min(1, relevance + 0.1),
                            'confidence_level': 0.95,
                            'factors': {'keyword_overlap': relevance}
                        })()
                        relevant_chunks.append(chunk)
                else:
                    # Normal threshold filtering
                    if relevance >= min_relevance:
                        chunk.relevance_score = type('obj', (object,), {
                            'score': relevance,
                            'confidence_lower': max(0, relevance - 0.1),
                            'confidence_upper': min(1, relevance + 0.1),
                            'confidence_level': 0.95,
                            'factors': {'keyword_overlap': relevance}
                        })()
                        relevant_chunks.append(chunk)
        
        # Sort by relevance
        relevant_chunks.sort(
            key=lambda c: c.relevance_score.score if c.relevance_score else 0.0,
            reverse=True
        )
        
        return relevant_chunks
    
    async def _validate_query(self, query: str) -> str:
        """Validate and clean the search query."""
        if not query or not query.strip():
            raise ValueError("Query cannot be empty")
        return query.strip()
    
    async def _update_stats(self, chunk_count: int) -> None:
        """Update processing statistics."""
        # This is a placeholder for statistics tracking
        pass
    
    async def _create_chunk(self, content: str, metadata: Dict[str, Any], token_count: int) -> ContextChunk:
        """Create a ContextChunk object."""
        from ..models import ContextChunk, ContextSource, SourceType
        return ContextChunk(
            content=content,
            source=ContextSource(
                name="document_source",
                source_type=SourceType.DOCUMENT,
                url=str(metadata.get('source_file', 'unknown'))
            ),
            metadata=metadata,
            token_count=token_count
        )
